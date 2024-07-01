from tkinter import *
from functools import partial
from tkinter import ttk
import tkinter as tk
import ttkthemes
from reportlab.pdfgen import canvas
from ttkthemes import ThemedStyle
import datetime
import pyautogui as pg
import time
from PIL import Image
import subprocess
import os
import mss
from pymongo.mongo_client import MongoClient
from pymongo.server_api import ServerApi
from bson import ObjectId
from docx import Document
from PIL import Image
from tkinter import Tk, Label, Button, filedialog
from pymongo import MongoClient
from PIL import Image
import io
from docx.shared import Inches
from tkinter import messagebox
import ssl
import certifi
# import win32api
# import win32print
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
ca = certifi.where()



# Universal resourse identification
uri = "mongodb+srv://mohitapte4:j3ZsXs6FGCnGATZm@cluster0.xmn1i2w.mongodb.net/?retryWrites=true&w=majority"

# Create a client object of class MongoClient
client = MongoClient(uri,tlsCAFile=ca)

# Function for User-Login validation
def validateLogin(username, password):    
    db = client['patient_data']
    collection = db['passwd']
    user = collection.find_one({"username": username.get(), "password": password.get()})
    if user:
        main_page()
    else:
       messagebox.showerror("Error", "Invalid Login Id or Password")


# Required lists for further project
x = []
medname = []
medtype = []
medadvice = []
days = []
dwm = []
qty = []
img_data = ""


def draw_multiline_text(canvas, text, x, y, width, height, font_size):
    lines = []
    current_line = ""
    words = text.split()
    max_line_height = 0

    for word in words:
        if canvas.stringWidth(current_line + " " + word, "Helvetica", font_size) < width:
            current_line += " " + word
        else:
            lines.append(current_line)
            current_line = word

    if current_line:
        lines.append(current_line)

    for line in lines:
        line_height = canvas.stringWidth(line, "Helvetica", font_size)
        if line_height > max_line_height:
            max_line_height = line_height

    total_lines = len(lines)
    remaining_height = height - (total_lines * max_line_height)
    y -= max_line_height

    for line in lines:
        canvas.drawString(x, y, line.strip())
        y -= max_line_height

    return remaining_height



class Patient():
    def __init__(self, mrd,fn,mn,ln,age,sex,address,mob,land,misc):
        self.mrd = mrd
        self.fn = fn
        self.mn = mn
        self.ln = ln
        self.age = age
        self.sex = sex
        self.address = address
        self.mob = mob
        self.land = land
        self.misc = misc
        




def new_patient():
    
    def validateSubmit(mrd,fn,mn,ln,age,sex,address,mob,land,misc):
        today = datetime.date.today()
        new_pat = Patient(mrd,fn,mn,ln,age,sex,address,mob,land,misc);
        if ((mob.get().isdigit()) and (land.get().isdigit()) and age.get().isdigit()) and fn.get().isalpha() and mn.get().isalpha() and ln.get().isalpha() and sex.get().isalpha():
            today_string = today.strftime('%d/%m/%Y')
            document = {"MRD":mrd.get(),"first_name":fn.get(),"middle_name":mn.get(),
                        "last_name":ln.get(),"age":age.get(),"sex":sex.get(),"address":address.get(),"mobile_no":mob.get(),
                        "land_no":land.get(), "misc":misc.get(),
                        'rds':'',
                        'rdc':'',
                        'rda':'',
                        'rdv':'',
                        'rcs':'',
                        'rcc':'',
                        'rca':'',
                        'rcv':'',
                        'rns':'',
                        'rnc':'',
                        'rna':'',
                        'rnv':'',
                        'lds':'',
                        'ldc':'',
                        'lda':'',
                        'ldv':'',
                        'lcs':'',
                        'lcc':'',
                        'lca':'',
                        'lcv':'',
                        'lns':'',
                        'lnc':'',
                        'lna':'',
                        'lnv':'',
                        'ipd':'',
                        'entry1':'',
                        'entry2':'',
                        'entry3':'',
                        'entry4':'',
                        'complaints':'Cheif Complaints:',
                        'examination':'Examination:',
                        'diagnosis':'Diagnosis:',
                        'medicine':'Medicine:',
                        'history':'History:',
                        'advised':'Advised:',
                        'x':[],
                        'medname':[],
                        'medtype':[],
                        'medadvice':[],
                        'days':[],
                        'dwm':[],
                        'qty':[],
                        'img_data':'',
                        'doatxt':'',
                        't1txt':'',
                        'dodtxt':'',
                        't2txt':'',
                        'cftxt':'',
                        'opnotestxt':'',
                        'investigationtxt':'',
                        'postmedicinetxt':'',
                        'surgeryadvisingtxt':'',
                        'adviseondischargetxt':'',
                        'date':''
                        }
            db = client.get_database('patient_data')
            collection = db['patient_name_age']
            collection.insert_one(document)
            
            cursor = collection.find({"$and": [
                                {"MRD": {"$regex": mrd.get(), "$options": "i"}},
                                {"first_name": {"$regex": fn.get(), "$options": "i"}},
                                {"middle_name": {"$regex": mn.get(), "$options": "i"}},
                                {"last_name": {"$regex": ln.get(), "$options": "i"}},
                                {"age": {"$regex": age.get(), "$options": "i"}},
                                {"address": {"$regex": address.get(), "$options": "i"}},
                                {"mobile_no": {"$regex": mob.get(), "$options": "i"}},
                                {"land_no": {"$regex": land.get(), "$options": "i"}},
                                {"misc": {"$regex": misc.get(), "$options": "i"}}
                            ]})


            data = [doc for doc in cursor]
            for doc in data:
                values = [str(v) for v in doc.values()]
            
                
            tab1.destroy()
            patient_selected(values)
        else:
            messagebox.showerror("Error", "An error occurred!")
            return
        
    tab1 = Toplevel(root)
    mrdLabel = ttk.Label(tab1, text="MRD Number").grid(row=0, column=0)
    mrd = StringVar()
    mrdEntry = ttk.Entry(tab1, textvariable=mrd).grid(row=0, column=1)
    import random
    today = datetime.date.today()
    new_day = today.day
    mrd_str = str(new_day)+str(random.randint(100000, 999999))
    mrd.set(mrd_str)
    
    
    fnLabel = ttk.Label(tab1, text="First Name").grid(row=2, column=0)
    fn = StringVar()
    fnEntry = ttk.Entry(tab1, textvariable=fn).grid(row=2, column=1) 

    mnLabel = ttk.Label(tab1, text="Middle Name").grid(row=4, column=0)
    mn = StringVar()
    mnEntry = ttk.Entry(tab1, textvariable=mn).grid(row=4, column=1) 

    lnLabel = ttk.Label(tab1, text="Last Name").grid(row=6, column=0)
    ln = StringVar()
    lnEntry = ttk.Entry(tab1, textvariable=ln).grid(row=6, column=1) 

    ageLabel = ttk.Label(tab1, text="Age").grid(row=8, column=0)
    age = StringVar()
    ageEntry = ttk.Entry(tab1, textvariable=age).grid(row=8, column=1) 

    sexLabel = ttk.Label(tab1, text="Sex").grid(row=10, column=0)
    sex = StringVar()
    ttk.Radiobutton(tab1,variable=sex, text="Male",value="Male", command=None).grid(row=10, column=1)
    ttk.Radiobutton(tab1,variable=sex, text="Female",value="Female", command=None).grid(row=10, column=2)

    addressLabel = ttk.Label(tab1, text="Address").grid(row=12, column=0)
    address = StringVar()
    addressEntry = ttk.Entry(tab1, textvariable=address).grid(row=12, column=1) 

    mobLabel = ttk.Label(tab1, text="Mobile Number").grid(row=14, column=0)
    mob = StringVar()
    mobEntry = ttk.Entry(tab1, textvariable=mob).grid(row=14, column=1) 

    landLabel = ttk.Label(tab1, text="Landline Number").grid(row=16, column=0)
    land = StringVar()
    landEntry = ttk.Entry(tab1, textvariable=land).grid(row=16, column=1)


    miscLabel = ttk.Label(tab1, text="Miscellaneous").grid(row=18, column=0)
    misc = StringVar()
    miscEntry = ttk.Entry(tab1, textvariable=misc).grid(row=18, column=1)

    validateSubmit = partial(validateSubmit, mrd,fn,mn,ln,age,sex,address,mob,land,misc)
    submitButton = ttk.Button(tab1, text="Submit", command=validateSubmit).grid(row=20, column=0)
    

def old_patient():
    
    def validateSubmit(mrd,fn,mn,ln,age,sex,address,mob,land,misc):

        
            def on_tree_select(event):
                # This line retrieves the ID of the selected item in the treeview widget that triggered the event.
                item = event.widget.selection()[0]
                # This line retrieves the values associated with the selected item in the treeview widget. It accesses the dictionary stored under the key 'values' within the dictionary returned by event.widget.item(item).
                values = event.widget.item(item)['values']
                #  This line destroys (closes) the tab2 frame. It assumes that tab2 is a Tkinter Frame object representing a tab in a notebook-style interface.
                tab2.destroy()
                # This line calls the patient_selected function and passes the values retrieved from the selected item as an argument. This function presumably takes some action based on the selected patient's information.
                patient_selected(values)
                


            # This line retrieves the database named 'patient_data' from the MongoDB client (client) 
            db = client.get_database('patient_data')

            #  This line retrieves the collection named 'patient_name_age' from the db database. It assumes that 'patient_name_age' is a collection (similar to a table in relational databases) within the 'patient_data' database.
            collection = db['patient_name_age']
            
            # this query retrieves documents from the MongoDB collection based on multiple conditions specified for different fields, and all conditions must be met for a document to be returned in the cursor.

            cursor = collection.find({"$and": [
                                {"MRD": {"$regex": mrd.get(), "$options": "i"}},   
                                {"first_name": {"$regex": fn.get(), "$options": "i"}},
                                {"middle_name": {"$regex": mn.get(), "$options": "i"}},
                                {"last_name": {"$regex": ln.get(), "$options": "i"}},
                                {"age": {"$regex": age.get(), "$options": "i"}},
                                {"sex": {"$regex": sex.get()}},
                                {"address": {"$regex": address.get(), "$options": "i"}},
                                {"mobile_no": {"$regex": mob.get(), "$options": "i"}},
                                {"land_no": {"$regex": land.get(), "$options": "i"}},
                                {"misc": {"$regex": misc.get(), "$options": "i"}}
                            ]})


            # This line creates a list called data using a list comprehension. It iterates over the cursor object and adds each document (doc) returned by the cursor to the list.
            data = [doc for doc in cursor]


            # This line creates a new Tkinter Frame widget named tree_frame. The tree_frame is a child of the tab2 widget (which presumably represents a tab in a notebook-style interface)
            tree_frame = ttk.Frame(tab2)
            # This line uses the grid geometry manager to place the tree_frame widget within the tab2 widget. It specifies that the tree_frame should be placed in the fourth column and thirtieth row of the grid, and it should expand in both the horizontal and vertical directions (sticky='nsew'). This allows the tree_frame to fill the available space within the tab2 widget.
            tree_frame.grid(column=4, row=30, sticky='nsew')
    


            # Create a tkinter window and Treeview widget to display the JSON data
            style = ttk.Style()
            style.configure("Custom.Treeview", rowheight=35)
            

            # This line creates a new Treeview widget named tree as a child of the tab2 widget. The columns parameter specifies the column names based on the keys of the first document in the data list. The show='headings' parameter specifies that only the column headings should be shown.
            tree = ttk.Treeview(tab2, columns=list(data[0].keys()), show='headings')
            
            # This line applies the custom style "Custom.Treeview" to the Treeview widget created earlier.
            tree.configure(style="Custom.Treeview")


            # This loop iterates over the keys of the first document in the data list (assuming there is at least one document in the list) and sets each key as the text for the corresponding column heading in the Treeview widget.
            for key in data[0].keys():
                tree.heading(key, text=key)



            # This loop iterates over each document in the data list. It converts the values of each document to strings and inserts them into the Treeview widget as a new row. It also checks if the value at index 1 (presumably representing the MRD) is not already in the mrd_data dictionary before inserting it.
            mrd_data={}
            for doc in data:
                values = [str(v) for v in doc.values()]
                if values[1] not in mrd_data:    
                    mrd_data[values[1]] = 1
                    tree.insert('', 'end', values=values)
                    
                    

            #This line creates a vertical scrollbar named tree_scroll for the Treeview widget. 
            tree_scroll = ttk.Scrollbar(tree_frame, orient='vertical', command=tree.yview)
            # This line configures the yscrollcommand of the Treeview widget to be controlled by the tree_scroll scrollbar.
            tree.configure(yscrollcommand=tree_scroll.set)
            # This line uses the grid geometry manager to position the Treeview widget within the tree_frame widget, which is a child of tab2. The sticky='nsew' option makes the widget expand in all directions.
            tree.grid(column=4, row=30, sticky='nsew')

            # This line positions the scrollbar tree_scroll adjacent to the Treeview widget in the vertical direction (sticky='ns').
            tree_scroll.grid(column=4, row=30, sticky='ns')

            # Bind the treeview event to a function that will be called when a row is selected
            tree.bind('<<TreeviewSelect>>', on_tree_select)
       
            

        

    
    tab2 = Toplevel(root)
    screen_width = tab2.winfo_screenwidth()
    screen_height = tab2.winfo_screenheight()
    tab2.geometry("%dx%d" % (screen_width, screen_height))
    tab2.title("Old Patient")
    mrdLabel = ttk.Label(tab2, text="MRD Number").grid(row=0, column=0)
    mrd = StringVar()
    mrdEntry = ttk.Entry(tab2, textvariable=mrd).grid(row=0, column=1)

    fnLabel = ttk.Label(tab2, text="First Name").grid(row=2, column=0)
    fn = StringVar()
    fnEntry = ttk.Entry(tab2, textvariable=fn).grid(row=2, column=1) 

    mnLabel = ttk.Label(tab2, text="Middle Name").grid(row=4, column=0)
    mn = StringVar()
    mnEntry = ttk.Entry(tab2, textvariable=mn).grid(row=4, column=1) 

    lnLabel = ttk.Label(tab2, text="Last Name").grid(row=6, column=0)
    ln = StringVar()
    lnEntry = ttk.Entry(tab2, textvariable=ln).grid(row=6, column=1) 

    ageLabel = ttk.Label(tab2, text="Age").grid(row=8, column=0)
    age = StringVar()
    ageEntry = ttk.Entry(tab2, textvariable=age).grid(row=8, column=1) 
    
    sexLabel = ttk.Label(tab2, text="Sex").grid(row=10, column=0)
    sex = StringVar()
    ttk.Radiobutton(tab2,variable=sex, text="Male",value="Male", command=None).grid(row=10, column=1)
    ttk.Radiobutton(tab2,variable=sex, text="Female",value="Female", command=None).grid(row=10, column=2)

    addressLabel = ttk.Label(tab2, text="Address").grid(row=12, column=0)
    address = StringVar()
    addressEntry = ttk.Entry(tab2, textvariable=address).grid(row=12, column=1) 

    mobLabel = ttk.Label(tab2, text="Mobile Number").grid(row=14, column=0)
    mob = StringVar()
    mobEntry = ttk.Entry(tab2, textvariable=mob).grid(row=14, column=1) 

    landLabel = ttk.Label(tab2, text="Landline Number").grid(row=16, column=0)
    land = StringVar()
    landEntry = ttk.Entry(tab2, textvariable=land).grid(row=16, column=1)


    miscLabel = ttk.Label(tab2, text="Miscellaneous").grid(row=18, column=0)
    misc = StringVar()
    miscEntry = ttk.Entry(tab2, textvariable=misc).grid(row=18, column=1)

    validateSubmit = partial(validateSubmit, mrd,fn,mn,ln,age,sex, address,mob,land,misc)
    submitButton = ttk.Button(tab2, text="Search", command=validateSubmit).grid(row=20, column=0)

   




def main_page():
    # Toplevel() is a class in Tkinter used to create independent windows (also known as "top-level windows" or "child windows") that are separate from the main application window (Tk() window). 
    app = Toplevel(root)

    # retrieve the width of the screen where the app window is currently located.
    screen_width = app.winfo_screenwidth()
    # retrieve the height in same way
    screen_height = app.winfo_screenheight()
    

    # sets the geometry of the app window to be equal to the width and height of the screen
    app.geometry("%dx%d" % (screen_width, screen_height))
    
    # creates a notebook-style tab control (ttk.Notebook) inside the app window. The tabControl variable is used to reference this tab control.
    tabControl = ttk.Notebook(app)
    
    # creates a frame (ttk.Frame) named tab1 to serve as the content of the first tab. This frame is added as a tab to the tabControl notebook.
    tab1 = ttk.Frame(tabControl)
    # tab 2
    tab2 = ttk.Frame(tabControl)

    # adds tab1 as a tab to the tabControl notebook, with the text label Out Patient Department
    tabControl.add(tab1, text ='Out Patient Department')
    # adds tab2 as a tab to the tabControl notebook, with the text label In Patient Department
    tabControl.add(tab2, text ='In Patient Department')


    # This packs the tabControl notebook inside the app window, causing it to expand to fill the available space in both the horizontal and vertical directions.
    tabControl.pack(expand = 1, fill ="both")
    
    # This creates a button (ttk.Button) inside tab1 with the specified text label "New Patient" and associates it with the "new_patient" function. The button is positioned using the grid geometry manager at row 1, column 0 within tab1.
    ttk.Button(tab1, text="New Patient", command=new_patient).grid(row=1, column=0) 
    # button in row 1 , col 1 -- "old_patient" function
    ttk.Button(tab1, text="Old Patient", command=old_patient).grid(row=1, column=1)
    


    global patient_selected
    global cur_pat


    def patient_selected(doc):
        global cur_pat
        # doc is list of values : Values from table , where we clicked and get into here
        # assign those values to new list called cur_pat
        cur_pat = doc
        # retrieve a database named patient_data using client object 
        db = client.get_database('patient_data')
        # select a collection named 'patient_name_age' and assign it to a variable 
        collection = db['patient_name_age']
        

        # here we will perform database query using pymongo library 
        # we will run a query to find a specific document in MongoDB collection 
        cursor = collection.find({"$and": [
                            {"first_name": {"$regex": str(doc[2]), "$options": "i"}},
                            {"middle_name": {"$regex": str(doc[3]), "$options": "i"}},
                            {"last_name": {"$regex": str(doc[4]), "$options": "i"}},
                        ]})
        # in above query we have used cursor object which can be used to iterate over retrieved documents 
        # construct a regex pattern using value present at doc[2] (option i makes regex case insensitive)
        # and the above regex with another regex which is created using value present at doc[3] and further and this with regex created using value present at doc[4]

        # anding them essentially created a complete regex = {first name , middle name, last name}
        # we then find it in collection , and return to cursor



        # lets say cursor is list and it will store the matched item from collection

        # let us store value present at cursor[0] in var document ( i can see from mongo db id that cursor[0] represents _id for selected person )

        document = cursor[0]

        # so the name that we will display on top can be created by concating doc[2] , doc[3] , doc[4] from table value lists

        name = doc[2] + ' ' + doc[3] + " " + doc[4]

        # get real time date
        today = datetime.date.today()
        # get real time 
        today_string = today.strftime('%d/%m/%Y')



        # creates a labeled frame titled "Patient Information" using the ttk.LabelFrame widget from the ttk module. It is placed inside the container tab1.
        patient_info_frame = ttk.LabelFrame(tab1, text = "Patient Information")
        # This line specifies the position of the patient_info_frame within its container using the grid geometry manager. It is placed in row 2 and column 0 of the parent widget (tab1).
        patient_info_frame.grid(row = 2, column = 0)
        

        # This line creates a label widget inside the patient_info_frame displaying the patient's name. The label's text is set to "Name: " followed by the value of the name variable. The label is placed in row 3 and column 0 within patient_info_frame.
        ttk.Label(patient_info_frame, text="Name: "+ name, borderwidth=3, relief="ridge").grid(row = 3, column= 0)
        ttk.Label(patient_info_frame, text="Age: " +str(doc[5]), borderwidth=3, relief="ridge").grid(row = 3, column= 1)
        ttk.Label(patient_info_frame, text="Sex: "+str(doc[6]), borderwidth=3, relief="ridge").grid(row = 3, column= 2)
        ttk.Label(patient_info_frame, text="Mob: "+str(doc[8]), borderwidth=3, relief="ridge").grid(row = 3, column= 3)
        ttk.Label(patient_info_frame, text="Date: "+today_string, borderwidth=3, relief="ridge").grid(row = 3, column= 4)
        


        # create a new frame , named as below , inside tab1 called patient details
        patient_detail_frame = ttk.LabelFrame(tab1, text = "Patient Details")
        patient_detail_frame.grid(row = 4, column = 0)




        def chief_complaints(event):
            complaints = Toplevel(root)
            complaints.geometry("500x500")
            
        


        def history_event(event):
            history = Toplevel(root)
            history.geometry("500x500")
        


        def exam(event):
            exam = Toplevel(root)
            exam.geometry("500x500")
            
        def diagram(event):
            diagram = Toplevel(root)
            diagram.geometry("500x500")
            global img_data

            def select_image():
                global img_data
                filename = filedialog.askopenfilename(initialdir="/", title="Select Image", filetypes=(("Image files", "*.jpg *.jpeg *.png"), ("All files", "*.*")))
                if filename:
                    image = Image.open(filename)
                    image.show()
            
                    # Store the image in MongoDB
                    with open(filename, "rb") as f:
                        img_data = f.read()
                        #print(img_data)
                    
                    
             
                    
            select_image()
           

                
                
            
            
            
            
        def diagnosis_event(event):
            diagnosis = Toplevel(root)
            diagnosis.geometry("500x500")
            



        def advised_event(event):
            advised = Toplevel(root)
            advised.geometry("500x500")
            
        
        
            
        
            
        def chief_medicine(event):
            # This line creates a new top-level window (a pop-up window) and assigns it to the variable medicine. The Toplevel function is used to create a new window, and root (presumably the main application window) is passed as the parent window, indicating that medicine is a child window of root
            medicine = Toplevel(root)
            # his line sets an attribute of the medicine window to make it display in fullscreen mode. The attributes method is used to set various attributes of the window, and "-fullscreen" is the attribute specifying fullscreen mode. The value True indicates that fullscreen mode is enabled.
            medicine.attributes("-fullscreen", True)

            screen_width = medicine.winfo_screenwidth()
            screen_height = medicine.winfo_screenheight()

            medicine.geometry("%dx%d" % (screen_width, screen_height))
            



            patient_info_frame = ttk.LabelFrame(medicine, text = "Patient Information")
            patient_info_frame.grid(row = 0, column = 0)
            


            medicine_frame = ttk.LabelFrame(medicine, text = "Medicine")
            medicine_frame.grid(row = 2, column = 0)



            ttk.Label(patient_info_frame, text="Name: "+ name, borderwidth=3, relief="ridge").grid(row = 9, column= 1)

            ttk.Label(patient_info_frame, text="Age: " +str(doc[5]), borderwidth=3, relief="ridge").grid(row = 9, column= 2)

            ttk.Label(patient_info_frame, text="Sex: "+str(doc[6]), borderwidth=3, relief="ridge").grid(row = 9, column= 3)

            ttk.Label(patient_info_frame, text="Mob: "+str(doc[8]), borderwidth=3, relief="ridge").grid(row = 9, column= 4)

            ttk.Label(patient_info_frame, text="Date: "+today_string+"             ", borderwidth=3, relief="ridge").grid(row = 9, column= 5)
            
            



            ttk.Label(medicine_frame, text="Root Of Administration").grid(row = 1, column= 1)
            ttk.Label(medicine_frame, text="Medicine Name").grid(row = 1, column= 5)
            ttk.Label(medicine_frame, text="Type").grid(row = 1, column= 6)
            ttk.Label(medicine_frame, text="Advice").grid(row = 1, column= 7)
            ttk.Label(medicine_frame, text="Days").grid(row = 1, column= 8)
            ttk.Label(medicine_frame, text="DWM").grid(row = 1, column= 9)
            ttk.Label(medicine_frame, text="Qty").grid(row = 1, column= 10)
            
            
            


            global currow
            currow = 1 
            vari = {}
            global x
            global medname
            global medtype
            global medadvice
            global days
            global dwm
            global qty
            
            
            
            def add_more():
                global currow
                currow+=1
                
                vari[str(currow)+"xtxt"] = ttk.Combobox(medicine_frame, values=["Right Eye", "Left Eye", "Both Eyes", "Oral", "IM", "IV"]) 
                vari[str(currow)+"xtxt"].grid(row=currow, column=1)
                
                

                
                vari[str(currow)+"mednametxt"] = Entry(medicine_frame, width=20)
                vari[str(currow)+"mednametxt"].grid(row=currow, column=5)
                
                
                vari[str(currow)+"typetxt"] = Entry(medicine_frame, width=4)
                vari[str(currow)+"typetxt"].grid(row=currow, column=6)
                
                vari[str(currow)+"medadvicetxt"] = Entry(medicine_frame, width=10)
                vari[str(currow)+"medadvicetxt"].grid(row=currow, column=7)
                
                
                vari[str(currow)+"daystxt"] = Entry(medicine_frame, width=4)
                vari[str(currow)+"daystxt"].grid(row=currow, column=8)
                
                
                vari[str(currow)+"dwmtxt"] = Entry(medicine_frame, width=2)
                vari[str(currow)+"dwmtxt"].grid(row=currow, column=9)
                
                
                vari[str(currow)+"qtytxt"] = Entry(medicine_frame, width=2)
                vari[str(currow)+"qtytxt"].grid(row=currow, column=10)
                
                
                try:    
                    vari[str(currow)+"xtxt"].insert(END, x[currow-2])
                    vari[str(currow)+"mednametxt"].insert(END, medname[currow-2])
                    vari[str(currow)+"typetxt"].insert(END, medtype[currow-2])
                    vari[str(currow)+"medadvicetxt"].insert(END, medadvice[currow-2])
                    vari[str(currow)+"daystxt"].insert(END, days[currow-2])
                    vari[str(currow)+"dwmtxt"].insert(END, dwm[currow-2])
                    vari[str(currow)+"qtytxt"].insert(END, qty[currow-2])
                except:
                    pass
     
        
                return
            
            
            for i in range(len(x)):
                add_more()
            
            
            
            def save():
                x.clear()

                medname.clear()
                medtype.clear()
                medadvice.clear()
                days.clear()
                dwm.clear()
                qty.clear()
                for i in range(currow-1):
                    x.append(vari[str(i+2)+"xtxt"].get())
                    medname.append(vari[str(i+2)+"mednametxt"].get())
                    medtype.append(vari[str(i+2)+"typetxt"].get())
                    medadvice.append(vari[str(i+2)+"medadvicetxt"].get())
                    days.append(vari[str(i+2)+"daystxt"].get())
                    dwm.append(vari[str(i+2)+"dwmtxt"].get())
                    qty.append(vari[str(i+2)+"qtytxt"].get())
                    
                print(x)   
                print(medname)
            
            


            # inside the mediciene window , create a button new 
            button = ttk.Button(medicine, text="New", command=add_more)
            button.grid(row = 12, column= 2, sticky=tk.S)
            
            # button save 
            button = ttk.Button(medicine, text="Save", command=save)
            button.grid(row = 12, column= 3, sticky=tk.S)
                
            





            def screenshot():
                pdf = canvas.Canvas("medicine.pdf")
                pdf.drawString(100, 800, name)
                pdf.drawString(250, 800, str(doc[5]))
                pdf.drawString(280, 800, str(doc[6]))
                pdf.drawString(330, 800, today_string)
                
                row = 750
                for i in range(currow - 1):
                
                    pdf.drawString(100, row, medname[i])
                    pdf.drawString(170, row, medtype[i])
                    pdf.drawString(270, row, medadvice[i])
                    pdf.drawString(370, row, "in "+x[i])
                    pdf.drawString(470, row, days[i]+" days")
                    pdf.drawString(520, row, "("+qty[i]+")")
                    
                    
                    pdf.drawString(100, row-30, "----------------------------------------------------------------------------------")
                    row -= 100
                pdf.save()
                
                filename = "medicine.pdf"
                
                if os.name == "posix":  # for macOS or Linux
                    os.system("open " + filename)
                elif os.name == "nt":  # for Windows
                    os.system("start " + filename)
            
            def exit_window():
                    medicine.destroy() 
            
            
            exit_button = ttk.Button(medicine, text="Exit", command=exit_window)
            exit_button.grid(row = 12, column= 5, sticky=tk.S)   
                
                
            button = ttk.Button(medicine_frame, text="Print", command=screenshot)
            button.grid(row = 12, column= 4, sticky=tk.S)
            
            
            
            
        def prescription(event):
            prescription = Toplevel(root)
            prescription.attributes("-fullscreen", True)
            screen_width = prescription.winfo_screenwidth()
            screen_height = prescription.winfo_screenheight()
            prescription.geometry("%dx%d" % (screen_width, screen_height))

            
            def insertValue(value, text_field, num_win):
                text_field.insert(END, value)
                num_win.destroy()
            
            def nums(event, field_txt):
                num_win = Toplevel(root)
                num_win.geometry("700x700")
                button1 = ttk.Button(num_win, text="0.75", command=lambda val=str(0.75): insertValue(val,field_txt,num_win))
                button1.grid(row=0, column=0)
                button2 = ttk.Button(num_win, text="0.1", command=lambda val=str(0.1): insertValue(val,field_txt,num_win))
                button2.grid(row=0, column=1)
                button3 = ttk.Button(num_win, text="0.2", command=lambda val=str(0.2): insertValue(val,field_txt,num_win))
                button3.grid(row=0, column=2)
                button4 = ttk.Button(num_win, text="0.25", command=lambda val=str(0.25): insertValue(val,field_txt,num_win))
                button4.grid(row=0, column=3)
                button5 = ttk.Button(num_win, text="0.3", command=lambda val=str(0.3): insertValue(val,field_txt,num_win))
                button5.grid(row=0, column=4)
                button6 = ttk.Button(num_win, text="0.35", command=lambda val=str(0.35): insertValue(val,field_txt,num_win))
                button6.grid(row=0, column=5)
                button7 = ttk.Button(num_win, text="0.4", command=lambda val=str(0.4): insertValue(val,field_txt,num_win))
                button7.grid(row=1, column=0)
                button8 = ttk.Button(num_win, text="0.45", command=lambda val=str(0.45): insertValue(val,field_txt,num_win))
                button8.grid(row=1, column=1)
                button9 = ttk.Button(num_win, text="0.5", command=lambda val=str(0.5): insertValue(val,field_txt,num_win))
                button9.grid(row=1, column=2)
                
            patient_info_frame = ttk.LabelFrame(prescription, text = "Patient Information")
            patient_info_frame.grid(row = 0, column = 0)
            
            glass_prescription_frame = ttk.LabelFrame(prescription, text = "Glass Prescription")
            glass_prescription_frame.grid(row = 2, column = 0)
            
            
            def exit():
                prescription.destroy()
            
            
            
            exit_button = ttk.Button(prescription, text="Exit", command=exit)
            exit_button.grid(row = 12, column= 5, sticky=tk.S)
            
            rds = document['rds']
            rdc = document['rdc']
            rda = document['rda']
            rdv = document['rdv']
            rcs = document['rcs']
            rcc = document['rcc']
            rca = document['rca']
            rcv = document['rcv']
            rns = document['rns']
            rnc = document['rnc']
            rna = document['rna']
            rnv = document['rnv']
            lds = document['lds']
            ldc = document['ldc']
            lda = document['lda']
            ldv = document['ldv']
            lcs = document['lcs']
            lcc = document['lcc']
            lca = document['lca']
            lcv = document['lcv']
            lns = document['lns']
            lnc = document['lnc']
            lna = document['lna']
            lnv = document['lnv']
            ipd = document['ipd']
            entry1txt = document['entry1']
            entry2txt = document['entry2']
            entry3txt = document['entry3']
            entry4txt = document['entry4']
              
            
            ttk.Label(patient_info_frame, text="Name: "+ name, borderwidth=3, relief="ridge").grid(row = 9, column= 1)
            ttk.Label(patient_info_frame, text="Age: " +str(doc[5]), borderwidth=3, relief="ridge").grid(row = 9, column= 2)
            ttk.Label(patient_info_frame, text="Sex: "+str(doc[6]), borderwidth=3, relief="ridge").grid(row = 9, column= 3)
            ttk.Label(patient_info_frame, text="Mob: "+str(doc[8]), borderwidth=3, relief="ridge").grid(row = 9, column= 4)
            ttk.Label(patient_info_frame, text="Date: "+today_string+"             ", borderwidth=3, relief="ridge").grid(row = 9, column= 5)
            
            
            
            
            
            ttk.Label(glass_prescription_frame, text="Right Eye").grid(row = 0, column= 1,columnspan=4)    
            ttk.Label(glass_prescription_frame, text="Left Eye").grid(row = 0, column= 5, columnspan=4)
            
            ttk.Label(glass_prescription_frame, text="Dist.").grid(row = 2, column= 0)
            ttk.Label(glass_prescription_frame, text="Computer").grid(row = 3, column= 0)
            ttk.Label(glass_prescription_frame, text="Near").grid(row = 4, column= 0)
            
            
            ttk.Label(glass_prescription_frame, text="Spl.").grid(row = 1, column= 1)
            ttk.Label(glass_prescription_frame, text="Cyl.").grid(row = 1, column= 2)
            ttk.Label(glass_prescription_frame, text="Axis").grid(row = 1, column= 3)
            ttk.Label(glass_prescription_frame, text="Vision").grid(row = 1, column= 4)
            
            ttk.Label(glass_prescription_frame, text="Spl.").grid(row = 1, column= 5)
            ttk.Label(glass_prescription_frame, text="Cyl.").grid(row = 1, column= 6)
            ttk.Label(glass_prescription_frame, text="Axis").grid(row = 1, column= 7)
            ttk.Label(glass_prescription_frame, text="Vision").grid(row = 1, column= 8)
            
            
            
            rdstxt = Entry(glass_prescription_frame, width=5)
            rdstxt.grid(row=2, column=1)
            rdstxt.bind("<Double-Button-1>", lambda event: nums(event, rdstxt), add="+")

            
            rdctxt = Entry(glass_prescription_frame, width=5)
            rdctxt.grid(row=2, column=2)
            rdctxt.bind("<Double-Button-1>", lambda event: nums(event, rdctxt), add="+")
            
            rdatxt = Entry(glass_prescription_frame, width=5)
            rdatxt.grid(row=2, column=3)
            rdatxt.bind("<Double-Button-1>", lambda event: nums(event, rdatxt), add="+")
            
            rdvtxt = Entry(glass_prescription_frame, width=5)
            rdvtxt.grid(row=2, column=4)
            rdvtxt.bind("<Double-Button-1>", lambda event: nums(event, rdvtxt), add="+")
            
            rcstxt = Entry(glass_prescription_frame, width=5)
            rcstxt.grid(row=3, column=1)
            rcstxt.bind("<Double-Button-1>", lambda event: nums(event, rcstxt), add="+")
            
            
            rcctxt = Entry(glass_prescription_frame, width=5)
            rcctxt.grid(row=3, column=2)
            rcctxt.bind("<Double-Button-1>", lambda event: nums(event, rcctxt), add="+")
            
            rcatxt = Entry(glass_prescription_frame, width=5)
            rcatxt.grid(row=3, column=3)
            rcatxt.bind("<Double-Button-1>", lambda event: nums(event, rcatxt), add="+")
            
            rcvtxt = Entry(glass_prescription_frame, width=5)
            rcvtxt.grid(row=3, column=4)
            rcvtxt.bind("<Double-Button-1>", lambda event: nums(event, rcvtxt), add="+")
            
            
            rnstxt = Entry(glass_prescription_frame, width=5)
            rnstxt.grid(row=4, column=1)
            rnstxt.bind("<Double-Button-1>", lambda event: nums(event, rnstxt), add="+")
            
            rnctxt = Entry(glass_prescription_frame, width=5)
            rnctxt.grid(row=4, column=2)
            rnctxt.bind("<Double-Button-1>", lambda event: nums(event, rnctxt), add="+")
            
            rnatxt = Entry(glass_prescription_frame, width=5)
            rnatxt.grid(row=4, column=3)
            rnatxt.bind("<Double-Button-1>", lambda event: nums(event, rnatxt), add="+")
            
            rnvtxt = Entry(glass_prescription_frame, width=5)
            rnvtxt.grid(row=4, column=4)
            rnvtxt.bind("<Double-Button-1>", lambda event: nums(event, rnvtxt), add="+")
            
            
            
            ldstxt = Entry(glass_prescription_frame, width=5)
            ldstxt.grid(row=2, column=5)
            ldstxt.bind("<Double-Button-1>", lambda event: nums(event, ldstxt), add="+")
            
            ldctxt = Entry(glass_prescription_frame, width=5)
            ldctxt.grid(row=2, column=6)
            ldctxt.bind("<Double-Button-1>", lambda event: nums(event, ldctxt), add="+")
            
            ldatxt = Entry(glass_prescription_frame, width=5)
            ldatxt.grid(row=2, column=7)
            ldatxt.bind("<Double-Button-1>", lambda event: nums(event, ldatxt), add="+")
            
            ldvtxt = Entry(glass_prescription_frame, width=5)
            ldvtxt.grid(row=2, column=8)
            ldvtxt.bind("<Double-Button-1>", lambda event: nums(event, ldvtxt), add="+")
            
            
            lcstxt = Entry(glass_prescription_frame, width=5)
            lcstxt.grid(row=3, column=5)
            lcstxt.bind("<Double-Button-1>", lambda event: nums(event, lcstxt), add="+")
            
            lcctxt = Entry(glass_prescription_frame, width=5)
            lcctxt.grid(row=3, column=6)
            lcctxt.bind("<Double-Button-1>", lambda event: nums(event, lcctxt), add="+")
            
            lcatxt = Entry(glass_prescription_frame, width=5)
            lcatxt.grid(row=3, column=7)
            lcatxt.bind("<Double-Button-1>", lambda event: nums(event, lcatxt), add="+")
            
            lcvtxt = Entry(glass_prescription_frame, width=5)
            lcvtxt.grid(row=3, column=8)
            lcvtxt.bind("<Double-Button-1>", lambda event: nums(event, lcvtxt), add="+")
            
            
            lnstxt = Entry(glass_prescription_frame, width=5)
            lnstxt.grid(row=4, column=5)
            lnstxt.bind("<Double-Button-1>", lambda event: nums(event, lnstxt), add="+")
            
            lnctxt = Entry(glass_prescription_frame, width=5)
            lnctxt.grid(row=4, column=6)
            lnctxt.bind("<Double-Button-1>", lambda event: nums(event, lnctxt), add="+")
            
            lnatxt = Entry(glass_prescription_frame, width=5)
            lnatxt.grid(row=4, column=7)
            lnatxt.bind("<Double-Button-1>", lambda event: nums(event, lnatxt), add="+")
            
            lnvtxt = Entry(glass_prescription_frame, width=5)
            lnvtxt.grid(row=4, column=8)
            lnvtxt.bind("<Double-Button-1>", lambda event: nums(event, lnvtxt), add="+")
            
            
            glass_details_frame = ttk.LabelFrame(prescription, text = "Glass Details")
            glass_details_frame.grid(row = 5, column = 0)
            
            ttk.Label(glass_details_frame, text="IPD:").grid(row = 5, column= 0)
            ipdtxt = Entry(glass_details_frame)
            ipdtxt.grid(row=5, column=1)
            ipdtxt.bind("<Double-Button-1>", lambda event: nums(event, ipdtxt), add="+")
            
            
            ttk.Label(glass_details_frame, text="Purpose: ").grid(row = 7, column= 0)
            global entry1
            entry1 = ttk.Combobox(glass_details_frame, values=["Constant Use", "For Near Vision Only", "For Distant Vision Only"])
            entry1.grid(row = 7, column= 1)
            entry2 = ttk.Combobox(glass_details_frame, values=["White Glass", "Photo Grey", "High Index Glass", "Anti-Glare Coating", "Goggles"])
            entry2.grid(row = 7, column= 2)
            
            
            ttk.Label(glass_details_frame, text="Remark: ").grid(row = 8, column= 0)
            entry3 = ttk.Combobox(glass_details_frame, values=["Kryptok Bifocal Glass", "Executive Bifocal Glass", "Progressive Glass"])
            entry3.grid(row = 8, column= 1)
            entry4 = ttk.Combobox(glass_details_frame, values=["Change Right Glass Only","Change Left Glass Only", "Change Both Glasses"])
            entry4.grid(row = 8, column= 2)
            
            
            rdstxt.insert(END, rds)
            rdctxt.insert(END, rdc)
            rdatxt.insert(END, rda)
            rdvtxt.insert(END, rdv)
            rcstxt.insert(END, rcs)
            rcctxt.insert(END, rcc)
            rcatxt.insert(END, rca)
            rcvtxt.insert(END, rcv)
            rnstxt.insert(END, rns)
            rnctxt.insert(END, rnc)
            rnatxt.insert(END, rna)
            rnvtxt.insert(END, rnv)
            ldstxt.insert(END, lds)
            ldctxt.insert(END, ldc)
            ldatxt.insert(END, lda)
            ldvtxt.insert(END, ldv)
            lcstxt.insert(END, lcs)
            lcctxt.insert(END, lcc)
            lcatxt.insert(END, lca)
            lcvtxt.insert(END, lcv)
            lnstxt.insert(END, lns)
            lnctxt.insert(END, lnc)
            lnatxt.insert(END, lna)
            lnvtxt.insert(END, lnv)
            ipdtxt.insert(END, ipd)
            entry1.insert(END, entry1txt)
            entry2.insert(END, entry2txt)
            entry3.insert(END, entry3txt)
            entry4.insert(END, entry4txt)
                        
            
            
            def save():
                filter = {"$and": [
                                    {"first_name": {"$regex": str(doc[2]), "$options": "i"}},
                                    {"middle_name": {"$regex": str(doc[3]), "$options": "i"}},
                                    {"last_name": {"$regex": str(doc[4]), "$options": "i"}},
                                ]}


                new_values = {'$set': {'rds': rdstxt.get(),
                       'rdc': rdctxt.get(),
                       'rda': rdatxt.get(),
                       'rdv': rdvtxt.get(),
                       'rcs': rcstxt.get(),
                       'rcc': rcctxt.get(),
                       'rca': rcatxt.get(),
                       'rcv': rcvtxt.get(),
                       'rns': rnstxt.get(),
                       'rnc': rnctxt.get(),
                       'rna': rnatxt.get(),
                       'rnv': rnvtxt.get(),
                       'lds': ldstxt.get(),
                       'ldc': ldctxt.get(),
                       'lda': ldatxt.get(),
                       'ldv': ldvtxt.get(),
                       'lcs': lcstxt.get(),
                       'lcc': lcctxt.get(),
                       'lca': lcatxt.get(),
                       'lcv': lcvtxt.get(),
                       'lns': lnstxt.get(),
                       'lnc': lnctxt.get(),
                       'lna': lnatxt.get(),
                       'lnv': lnvtxt.get(),
                       'ipd': ipdtxt.get(),
                       'entry1': entry1.get(),
                       'entry2': entry2.get(),
                       'entry3': entry3.get(),
                       'entry4': entry4.get()}}


                result = collection.update_one(filter, new_values)
            
            
            button = ttk.Button(prescription, text="              Save", command=save)
            button.grid(row = 12, column= 4, sticky=tk.S)
            
            
            
            
            
            def screenshot():
                filename = "screenshot.png"
                
                cropped_image_path = 'cropped_screenshot.png'
                
               
                print_command = ["lp", "-d", "printer-name", cropped_image_path]
                
                
                subprocess.run(print_command)
                
                subprocess.run(["xdg-open", cropped_image_path])
                
            

                
                
                
            button = ttk.Button(prescription, text="              Print", command=screenshot)
            button.grid(row = 12, column= 3, sticky=tk.S)
            
        complaints = document['complaints']
        examination = document['examination']
        diagnosis = document['diagnosis']
        medicine = document['medicine']
        history = document['history']
        advised = document['advised'] 
            












        ttk.Label(patient_detail_frame, text="COMPLAINTS").grid(row=3, column=0)
        # This line creates a text widget named complaintxt within the patient_detail_frame. The text widget is configured to have a height of 10 lines, a width of 25 characters, and a light yellow background color.
        complaintxt = Text(patient_detail_frame, height = 10,
                        width = 25,
                        bg = "light yellow")
        # This line specifies the position of the complaintxt text widget within the patient_detail_frame using the grid geometry manager. It is placed in row 4 and column 0.
        complaintxt.grid(row=4, column=0)
        # This line inserts the string "Complaints" at the end (i.e., after any existing text) of the complaintxt text widget.
        complaintxt.insert(END, "Complaints")

        # This line binds the <Double-Button-1> event (double-click of the left mouse button) to the chief_complaints function when it occurs within the complaintxt text widget. The add="+" argument ensures that this binding does not replace any existing bindings for the same event.
        complaintxt.bind("<Double-Button-1>", chief_complaints, add="+")
        
        
        
        


        
        
        
        ttk.Label(patient_detail_frame, text="HISTORY").grid(row=5, column=0)
        historytxt = Text(patient_detail_frame, height = 10,
                        width = 25,
                        bg = "light yellow")
        
        historytxt.grid(row=6, column=0)
        
        historytxt.insert(END, "History")

        historytxt.bind("<Double-Button-1>", history_event, add="+")
        
        




        
        ttk.Label(patient_detail_frame, text="EXAMINATION").grid(row=3, column=1)
        examtxt = Text(patient_detail_frame, height = 10,
                        width = 25,
                        bg = "light yellow")
        
        examtxt.grid(row=4, column=1)
        
        examtxt.insert(END, "Examination")
        examtxt.bind("<Double-Button-1>", exam, add="+")
        




        
        ttk.Label(patient_detail_frame, text="DIAGRAM").grid(row=5, column=1)
        diagramtxt = Text(patient_detail_frame, height = 10,
                        width = 25,
                        bg = "light yellow")
        
        diagramtxt.grid(row=6, column=1)
        
        diagramtxt.insert(END, "Diagram")
        diagramtxt.bind("<Double-Button-1>", diagram, add="+")
            
        




        ttk.Label(patient_detail_frame, text="DIAGNOSIS").grid(row=3, column=2)
        diagnosistxt = Text(patient_detail_frame, height = 10,
                        width = 25,
                        bg = "light yellow")
        
        diagnosistxt.grid(row=4, column=2)
        
        diagnosistxt.insert(END, "Diagnosis")
        diagnosistxt.bind("<Double-Button-1>", diagnosis_event, add="+")
        
        






        ttk.Label(patient_detail_frame, text="ADVISED").grid(row=5, column=2)
        advisedtxt = Text(patient_detail_frame, height = 10,
                        width = 25,
                        bg = "light yellow")
        
        advisedtxt.grid(row=6, column=2)
        
        advisedtxt.insert(END, "Advised")
        advisedtxt.bind("<Double-Button-1>", advised_event, add="+")
        







        
        ttk.Label(patient_detail_frame, text="MEDICINE").grid(row=3, column=3)
        medicinetxt = Text(patient_detail_frame, height = 10,
                        width = 25,
                        bg = "light yellow")
        
        medicinetxt.grid(row=4, column=3)
        
        medicinetxt.insert(END, "Medicine")
        medicinetxt.bind("<Double-Button-1>", chief_medicine, add="+")
        
        





        ttk.Label(patient_detail_frame, text="PRESCRIPTION").grid(row=5, column=3)
        prescriptiontxt = Text(patient_detail_frame, height = 10,
                        width = 25,
                        bg = "light yellow")
        
        prescriptiontxt.grid(row=6, column=3)
        
        prescriptiontxt.insert(END, "Prescription:")
        prescriptiontxt.bind("<Double-Button-1>", prescription, add="+")
        
        def new_save():
            filter = {"$and": [
                                {"first_name": {"$regex": str(doc[2]), "$options": "i"}},
                                {"middle_name": {"$regex": str(doc[3]), "$options": "i"}},
                                {"last_name": {"$regex": str(doc[4]), "$options": "i"}},
                            ]}


            global img_data
            result = collection.find_one(filter)
            source_dict = dict(result)
            new_id = ObjectId()
            today_string_with_timestamp = datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S')
            source_dict['complaints'] = complaintxt.get("1.0",'end-1c')
            source_dict['examination'] = examtxt.get("1.0",'end-1c')
            source_dict['diagnosis'] = diagnosistxt.get("1.0",'end-1c')
            source_dict['medicine'] = medicinetxt.get("1.0",'end-1c')
            source_dict['history'] = historytxt.get("1.0",'end-1c')
            source_dict['advised'] = advisedtxt.get("1.0",'end-1c')
            source_dict['date'] = today_string_with_timestamp
            source_dict['_id'] = new_id
            
            
            
            source_dict['doatxt'] = doatxt.get()
            source_dict['t1txt'] = t1txt.get()
            source_dict['dodtxt'] = dodtxt.get()
            source_dict['t2txt'] = t2txt.get()
            source_dict['cftxt'] = cftxt.get("1.0",'end-1c')
            source_dict['opnotestxt'] = opnotestxt.get("1.0",'end-1c')
            source_dict['investigationtxt'] = investigationtxt.get("1.0",'end-1c')
            source_dict['postmedicinetxt'] = postmedicinetxt.get("1.0",'end-1c')
            source_dict['surgeryadvisingtxt'] = surgeryadvisingtxt.get("1.0",'end-1c')
            source_dict['adviseondischargetxt'] = adviseondischargetxt.get("1.0",'end-1c')
            
            
            
            
            
            
            source_dict['x'] = x
            source_dict['medname'] = medname
            source_dict['medtype'] = medtype
            source_dict['medadvice'] = medadvice
            source_dict['days'] = days
            source_dict['dwm'] = dwm
            source_dict['qty'] = qty
            source_dict['img_data'] = img_data
            print(img_data)            
            collection.insert_one(source_dict)
        
        
        new_button = ttk.Button(patient_detail_frame, text="Save", command=new_save)
        new_button.grid(row = 12, column= 4, sticky=tk.S)
        
        
        def view_history():
            view_hist = Toplevel(root)
            view_hist.geometry("1000x1000")
            global cur_pat
            
            
            def retrive_data(date):
                db = client.get_database('patient_data')
                collection = db['patient_name_age']
                new_col = collection.find({"$and": [
                                    {"first_name": {"$regex": str(cur_pat[2]), "$options": "i"}},
                                    {"middle_name": {"$regex": str(cur_pat[3]), "$options": "i"}},
                                    {"last_name": {"$regex": str(cur_pat[4]), "$options": "i"}},
                                    {"date": {"$regex": str(date), "$options": "i"}}
                                ]})
                
                
                
                for document in new_col:

                    string = ""
                    string += str(document['complaints'])
                    string+="\n\n"
                    string += str(document['history'])
                    string+="\n\n"
                    string += str(document['examination'])
                    string+="\n\n"
                    string += str(document['diagnosis'])
                    string+="\n\n"
                    string += str(document['advised'])
                    string+="\n\n"
                    
                    for i in range(len(document['x'])):
                        
                        string += document['medname'][i] + "            "
                        string += document['medtype'][i] + "            "
                        string += document['medadvice'][i] + "            "
                        string += "in " + document['x'][i] + "         "
                        string += document['days'][i]+" days" + "            "
                        string += "("+document['qty'][i]+")"
                        string += "\n"
                        string += "--------------------------------"
                        string += "\n"
                    
                    
                    output_path = "output.docx"

                    doc = Document()
                    doc.add_paragraph(string)
                    
                        
                    try:
                        bin_data  = document['img_data']
                        print(bin_data)
                        image_stream = io.BytesIO(bin_data)
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run()
                        
                        # Insert the image into the run
                        run.add_picture(image_stream, width=Inches(4))  
                    except:
                        pass
                    
                    doc.save(output_path)

                    filename = "output.docx"
                    
                    if os.name == "posix":  # for macOS or Linux
                        os.system("open " + filename)
                    elif os.name == "nt":  # for Windows
                        os.system("start " + filename)
                

                
            db = client.get_database('patient_data')
            collection = db['patient_name_age']
            
            cursor = collection.find({"$and": [
                                {"first_name": {"$regex": str(cur_pat[2]), "$options": "i"}},
                                {"middle_name": {"$regex": str(cur_pat[3]), "$options": "i"}},
                                {"last_name": {"$regex": str(cur_pat[4]), "$options": "i"}},
                            ]})
            data = [doc for doc in cursor]
            new_dict = {}
            row = 0
            
            for doc in data:
                values = [str(v) for v in doc.values()]
                ttk.Button(view_hist, text=str(values[-1]), command=lambda date=values[-1]: retrive_data(date)).grid(row = row, column= 1, sticky=tk.S)
                row += 1
                
                    
          
            
            
        
        hist_button = ttk.Button(patient_detail_frame, text="Complete History", command=view_history)
        hist_button.grid(row = 13, column= 4, sticky=tk.S)
        
        
        


        def billing_section():
            billing = Toplevel(root)
            billing.attributes("-fullscreen", True)

            screen_width = billing.winfo_screenwidth()
            screen_height = billing.winfo_screenheight()

            billing.geometry("%dx%d" % (screen_width, screen_height))

            def calculate_amount():
                try:
                    unit = float(unit_entry.get())
                    rate = float(rate_entry.get())
                    discount = float(discount_entry.get() or 0)
                    amount = unit * rate * (1 - discount / 100)
                    total_amount_label.config(text=f"Total amount to be paid: {amount:.2f} Rs.")
                except:
                    #total_amount_label.config(text="Invalid input")
                    messagebox.showerror("Error", "An error occurred!")
                
                
            def convert_to_words(num):
                if num == 0:
                    return "zero"
                
                ones = ["", "one", "two", "three", "four", "five", "six", "seven", "eight", "nine"]
                tens = ["", "", "twenty", "thirty", "forty", "fifty", "sixty", "seventy", "eighty", "ninety"]
                teens = ["ten", "eleven", "twelve", "thirteen", "fourteen", "fifteen", "sixteen", "seventeen", "eighteen", "nineteen"]
                
                words = ""
                
                # Handling lakhs
                if num >= 100000:
                    lakhs = num // 100000
                    words += ones[lakhs] + " lakh "
                    num %= 100000
                
                # Handling thousands
                if num >= 1000:
                    thousands = num // 1000
                    words += convert_to_words(thousands) + " thousand "
                    num %= 1000
                
                # Handling hundreds
                if num >= 100:
                    hundreds = num // 100
                    words += ones[hundreds] + " hundred "
                    num %= 100
                    
                # Handling tens and ones
                if num >= 10 and num <= 19:
                    words += teens[num - 10] + " "
                    num = 0
                elif num >= 20:
                    tens_digit = num // 10
                    words += tens[tens_digit] + " "
                    num %= 10
                    
                if num >= 1 and num <= 9:
                    words += ones[num] + " "
                    
                return words.strip()

            def generate_pdf_bill(patient_info, billing_details, doctor_info, total_amount, department, billing_items):
                c = canvas.Canvas("OPD_bill.pdf", pagesize=letter)

                bold_font = "Helvetica-Bold"
                normal_font = "Helvetica"
                title_font_size = 15
                section_font_size = 13
                content_font_size = 12

                text_x = 50
                text_y = 720
                line_height = 20

                c.setFont(bold_font, title_font_size)
                c.drawCentredString(300, 760, "EYE HOSPITAL, PUNE")
                c.line(50, 740, 550, 740)

                
                c.setFont(bold_font, section_font_size)
                c.drawString(text_x, text_y, "Patient Information:")
                text_y -= line_height
                for label, value in patient_info.items():
                    c.setFont(normal_font, content_font_size)
                    c.drawString(text_x, text_y, f"{label}: {value}")
                    text_y -= line_height

                text_y -= line_height

                c.setFont(bold_font, section_font_size)
                c.drawString(text_x, text_y, "Billing Details:")
                text_y -= line_height
                for label, value in billing_details.items():
                    c.setFont(normal_font, content_font_size)
                    c.drawString(text_x, text_y, f"{label}: {value}")
                    text_y -= line_height

                # Add spacing between sections
                text_y -= line_height

                # Draw doctor information
                c.setFont(bold_font, section_font_size)
                c.drawString(text_x, text_y, "Doctor Information:")
                text_y -= line_height
                for label, value in doctor_info.items():
                    c.setFont(normal_font, content_font_size)
                    c.drawString(text_x, text_y, f"{label}: {value}")
                    text_y -= line_height

                # Add spacing between sections
                text_y -= line_height

                # Draw department
                c.setFont(bold_font, section_font_size)
                c.drawString(text_x, text_y, f"Department: {department}")
                text_y -= line_height

                # Draw billing items in a tabular format
                
                text_y -= line_height

                
                
                
                
               
                c.setFont(bold_font, section_font_size)
                c.drawString(text_x, text_y, f"Total Amount: Rs. {total_amount}")

        
                text_y -= line_height

                c.setFont(bold_font, section_font_size)
                c.drawString(text_x, text_y, "RECEIPT")
                text_y -= line_height
                c.setFont(normal_font, content_font_size)
                c.drawString(text_x, text_y, f"Received the sum of Rs {total_amount}")
                text_y -= line_height
                amt = int(total_amount)
                result = convert_to_words(amt)
                c.drawString(text_x, text_y, f"In words : {result}")
                text_y -= line_height
                c.drawString(text_x, text_y, "By CASH / CARD / CHEQUE / UPI")

                # Save the PDF
                c.save()
                filename = "OPD_bill.pdf"
                
                if os.name == "posix": 
                    os.system("open " + filename)
                elif os.name == "nt": 
                    os.system("start " + filename)



             # Patient Information Frame
            
            
            


            def print_receipt():
                # Gather patient, billing, and doctor information
                patient_info = {
                    "Name": name,  # This seems incorrect, it should be the patient's name
                    "Age": str(doc[5]),
                    "Sex": str(doc[6]),
                    "Mobile": str(doc[8]),
                    "Date": today_string
                }

                billing_details = {
                    "Billing Item": billing_item_entry.get(),  # Update to match the label in the GUI
                    "Unit": unit_entry.get(),
                    "Rate": rate_entry.get(),
                    "Discount": discount_entry.get(),
                    "Detail": detail_entry.get()
                }

                doctor_info = {
                    "Doctor's Name": doctor_name_entry.get(),
                    "Doctor's Degree": doctor_degree_entry.get()
                }
                
                # Gather additional information
                department = department_entry.get()
                billing_items = {
                        "name": billing_item_entry.get(),
                        "unit": unit_entry.get(),
                        "rate": rate_entry.get(),
                        "discount": discount_entry.get(),
                        "amount": "calculate amount here",  # You need to calculate this value
                        "detail": detail_entry.get()
                    }
                

                total_amount_text = total_amount_label.cget("text").split(":")[-1].strip().replace('Rs.', '').strip()
                try:
                    # Convert the extracted numeric part to a floating-point number
                    total_amount_float = float(total_amount_text)

                    # Convert the floating-point number to an integer (if needed) or use it directly
                    total_amount = int(total_amount_float)   
                except:
                    messagebox.showerror("Error", "An error occurred!")
                # Generate PDF receipt
                generate_pdf_bill(patient_info, billing_details, doctor_info, total_amount, department, billing_items)


            
            
            
            
            
            
            patient_info_frame = ttk.LabelFrame(billing, text="Patient Information")
            patient_info_frame.grid(row=0, column=0, padx=50, pady=20, sticky="nsew")  

            ttk.Label(patient_info_frame, text="Bill No: ").grid(row=0, column=0, sticky="w")
            bill_no_entry = ttk.Entry(patient_info_frame)
            bill_no_entry.grid(row=0, column=0, sticky="e")


            
            
            ttk.Label(patient_info_frame, text="Name of Patient: " + name).grid(row=1, column=0, sticky="w")

            
            ttk.Label(patient_info_frame, text="Age: " + str(doc[5])).grid(row=2, column=0, sticky="w")

            
            ttk.Label(patient_info_frame, text="Sex: " + str(doc[6])).grid(row=3, column=0, sticky="w")

            
            ttk.Label(patient_info_frame, text="Mobile Number: " + str(doc[8])).grid(row=4, column=0, sticky="w")

            
            ttk.Label(patient_info_frame, text="Date: " + today_string).grid(row=5, column=0, sticky="w")







            
            billing_details_frame = ttk.LabelFrame(billing, text="Billing Details")
            billing_details_frame.grid(row=1, column=0, padx=50, pady=20, sticky="nsew") 

            


        
            
            department_label = ttk.Label(billing_details_frame, text="Department: ")
            department_label.grid(row=0, column=0, sticky="e")

            department_entry = ttk.Combobox(billing_details_frame, values=["OPD", "IPD"])
            department_entry.grid(row=0, column=1, sticky="w")



            
            billing_item_label = ttk.Label(billing_details_frame, text="Billing Item: ")
            billing_item_label.grid(row=1, column=0, sticky="e")
            billing_item_entry = ttk.Entry(billing_details_frame)
            billing_item_entry.grid(row=1, column=1, sticky="w")

            
            unit_label = ttk.Label(billing_details_frame, text="Unit: ")
            unit_label.grid(row=2, column=0, sticky="e")
            unit_entry = ttk.Entry(billing_details_frame)
            unit_entry.grid(row=2, column=1, sticky="w")

            
            rate_label = ttk.Label(billing_details_frame, text="Rate: ")
            rate_label.grid(row=3, column=0, sticky="e")
            rate_entry = ttk.Entry(billing_details_frame)
            rate_entry.grid(row=3, column=1, sticky="w")

           
            discount_label = ttk.Label(billing_details_frame, text="Discount: ")
            discount_label.grid(row=4, column=0, sticky="e")
            discount_entry = ttk.Entry(billing_details_frame)
            discount_entry.grid(row=4, column=1, sticky="w")


            



            
            detail_label = ttk.Label(billing_details_frame, text="Detail: ")
            detail_label.grid(row=6, column=0, sticky="e")
            detail_entry = ttk.Entry(billing_details_frame)
            detail_entry.grid(row=6, column=1, sticky="w")

            
            total_amount_label = ttk.Label(billing_details_frame, text="Total amount to be paid: ")
            total_amount_label.grid(row=5, column=0, sticky="e")

            calculate_amount_button = ttk.Button(billing_details_frame, text="Calculate Amount", command=calculate_amount)
            calculate_amount_button.grid(row=7, columnspan=2, pady=10)

            
            

            ttk.Button(billing, text="Print Receipt", command=print_receipt).grid(row=4, column=0, padx=50, pady=10, sticky="nsew")

            ttk.Button(billing, text="Exit Billing Section", command=billing.destroy).grid(row=5, column=0, padx=50, pady=10, sticky="nsew")

            




            doctor_details_frame = ttk.LabelFrame(billing, text="Doctor Details")
            doctor_details_frame.grid(row=0, column=1, padx=50, pady=20, sticky="nsew")



            
            doctor_name_label = ttk.Label(doctor_details_frame, text="Doctor's Name: ")
            doctor_name_label.grid(row=0, column=0, sticky="e")
            doctor_name_entry = ttk.Combobox(doctor_details_frame, values=["Dr. Salim Pathan"])
            doctor_name_entry.grid(row=0, column=1, sticky="w")


            doctor_degree_label = ttk.Label(doctor_details_frame, text="Doctor's Degree: ")
            doctor_degree_label.grid(row=1, column=0, sticky="e")
            doctor_degree_entry = ttk.Combobox(doctor_details_frame, values=["MBBS", "MS", "MD"])
            doctor_degree_entry.grid(row=1, column=1, sticky="w")


       



            


        billing_button=ttk.Button(patient_detail_frame,text="Billing",command=billing_section)
        # billing section func to be implemented
        billing_button.grid(row=13 , column= 3,sticky=tk.S)


















        #Tab 2 IPD
        
        patient_info_ipd = ttk.LabelFrame(tab2, text = "Patient Information")
        patient_info_ipd.grid(row = 1, column = 0)
        patient_discharge_ipd = ttk.LabelFrame(tab2, text = "Discharge Summary")
        patient_discharge_ipd.grid(row = 2, column = 0)
        
        ttk.Label(patient_discharge_ipd, text="Date of Admission", borderwidth=3, relief="ridge").grid(row = 1, column= 0)
        ttk.Label(patient_discharge_ipd, text="Time", borderwidth=3, relief="ridge").grid(row = 1, column= 4)
        ttk.Label(patient_discharge_ipd, text="Date of Discharge", borderwidth=3, relief="ridge").grid(row = 2, column= 0)
        ttk.Label(patient_discharge_ipd, text="Time", borderwidth=3, relief="ridge").grid(row = 2, column= 4)
        
        doatxt = Entry(patient_discharge_ipd, width=10)
        doatxt.grid(row=1, column=1)
        
        t1txt = Entry(patient_discharge_ipd, width=10)
        t1txt.grid(row=1, column=5)
        
        
        dodtxt = Entry(patient_discharge_ipd, width=10)
        dodtxt.grid(row=2, column=1)
        
        t2txt = Entry(patient_discharge_ipd, width=10)
        t2txt.grid(row=2, column=5)
        
        
        ttk.Label(patient_info_ipd, text="Name: "+ name, borderwidth=3, relief="ridge").grid(row = 3, column= 0)
        ttk.Label(patient_info_ipd, text="Age: " +str(doc[5]), borderwidth=3, relief="ridge").grid(row = 3, column= 1)
        ttk.Label(patient_info_ipd, text="Sex: "+str(doc[6]), borderwidth=3, relief="ridge").grid(row = 3, column= 2)
        ttk.Label(patient_info_ipd, text="Mob: "+str(doc[8]), borderwidth=3, relief="ridge").grid(row = 3, column= 3)
        ttk.Label(patient_info_ipd, text="Date: "+today_string, borderwidth=3, relief="ridge").grid(row = 3, column= 4)
        
        patient_detail_ipd = ttk.LabelFrame(tab2, text = "Patient Details IPD")
        patient_detail_ipd.grid(row = 4, column = 0)
        
        
        ttk.Label(patient_detail_ipd, text="Clinical Findings").grid(row=3, column=0)
        cftxt = Text(patient_detail_ipd, height = 10,
                        width = 25,
                        bg = "light yellow")
        
        cftxt.grid(row=4, column=0)
        
        cftxt.insert(END, "Clinical Findings")
        
        
        ttk.Label(patient_detail_ipd, text="Operation Notes").grid(row=5, column=0)
        opnotestxt = Text(patient_detail_ipd, height = 10,
                        width = 25,
                        bg = "light yellow")
        
        opnotestxt.grid(row=6, column=0)
        
        opnotestxt.insert(END, "Operation Notes")
        
        
        
        ttk.Label(patient_detail_ipd, text="Invesitgation").grid(row=3, column=1)
        investigationtxt = Text(patient_detail_ipd, height = 10,
                        width = 25,
                        bg = "light yellow")
        
        investigationtxt.grid(row=4, column=1)
        
        investigationtxt.insert(END, "Invesitgation")
        
        
        ttk.Label(patient_detail_ipd, text="Post Operative Medicines").grid(row=5, column=1)
        postmedicinetxt = Text(patient_detail_ipd, height = 10,
                        width = 25,
                        bg = "light yellow")
        
        postmedicinetxt.grid(row=6, column=1)
        
        postmedicinetxt.insert(END, "Post Operative Medicines")

        

        ttk.Label(patient_detail_ipd, text="Surgery Advising").grid(row=3, column=2)
        surgeryadvisingtxt = Text(patient_detail_ipd, height = 10,
                        width = 25,
                        bg = "light yellow")
        
        surgeryadvisingtxt.grid(row=4, column=2)
        
        surgeryadvisingtxt.insert(END, "Surgery Advising")
        
        ttk.Label(patient_detail_ipd, text="Advice on Discharge").grid(row=5, column=2)
        adviseondischargetxt = Text(patient_detail_ipd, height = 10,
                        width = 25,
                        bg = "light yellow")
        
        adviseondischargetxt.grid(row=6, column=2)
        
        adviseondischargetxt.insert(END, "Advice on Discharge")
        
        ipd_button = ttk.Button(patient_detail_ipd, text="Save", command=new_save)
        ipd_button.grid(row = 12, column= 4, sticky=tk.S)
        
        
        def view_ipd_history():
            view_hist = Toplevel(root)
            view_hist.geometry("1000x1000")
            global cur_pat
            
            
            def retrive_data(date):
                db = client.get_database('patient_data')
                collection = db['patient_name_age']
                new_col = collection.find({"$and": [
                                    {"first_name": {"$regex": str(cur_pat[2]), "$options": "i"}},
                                    {"middle_name": {"$regex": str(cur_pat[3]), "$options": "i"}},
                                    {"last_name": {"$regex": str(cur_pat[4]), "$options": "i"}},
                                    {"date": {"$regex": str(date), "$options": "i"}}
                                ]})
                
                
                
                for document in new_col:

                    string = ""
                    
                    string += str(document['doatxt']) + "  "

                    
                    string += str(document['t1txt'])
                    string+="\n\n"
                    
                    string += str(document['dodtxt']) + "  "

                    
                    string += str(document['t2txt'])
                    string+="\n\n"
                    
                    
                    string += str(document['cftxt'])
                    string+="\n\n"
                    string += str(document['opnotestxt'])
                    string+="\n\n"
                    string += str(document['investigationtxt'])
                    string+="\n\n"
                    string += str(document['postmedicinetxt'])
                    string+="\n\n"
                    string += str(document['surgeryadvisingtxt'])
                    string+="\n\n"
                    string += str(document['adviseondischargetxt'])
                    string+="\n\n"
                    
                    output_path = "ipd.docx"

                    doc = Document()
                    doc.add_paragraph(string)
                    
                    
                    doc.save(output_path)

                    filename = "ipd.docx"
                    
                    if os.name == "posix":  # for macOS or Linux
                        os.system("open " + filename)
                    elif os.name == "nt":  # for Windows
                        os.system("start " + filename)
                

                
            db = client.get_database('patient_data')
            collection = db['patient_name_age']
            
            cursor = collection.find({"$and": [
                                {"first_name": {"$regex": str(cur_pat[2]), "$options": "i"}},
                                {"middle_name": {"$regex": str(cur_pat[3]), "$options": "i"}},
                                {"last_name": {"$regex": str(cur_pat[4]), "$options": "i"}},
                            ]})
            data = [doc for doc in cursor]
            new_dict = {}
            row = 0
            
            for doc in data:
                values = [str(v) for v in doc.values()]
                ttk.Button(view_hist, text=str(values[-1]), command=lambda date=values[-1]: retrive_data(date)).grid(row = row, column= 1, sticky=tk.S)
                row += 1
                
                    
          
            
            
        
        ipd_hist_button = ttk.Button(patient_detail_ipd, text="IPD History", command=view_ipd_history)
        ipd_hist_button.grid(row = 13, column= 4, sticky=tk.S)
        
        
        


        
        def ipd_billing_section():
            billing = Toplevel(root)
            billing.attributes("-fullscreen", False)

            screen_width = billing.winfo_screenwidth()
            screen_height = billing.winfo_screenheight()

            billing.geometry("%dx%d" % (screen_width, screen_height))

            def calculate_amount():
                try:
                    operative_fees = float(operative_fees_entry.get())
                    lens_implant = float(lens_implant_entry.get())
                    operation_theatre_charges = float(operation_theatre_charges_entry.get())
                    anaesthetist_fees = float(anaesthetist_fees_entry.get())
                    viscoelastic = float(viscoelastic_entry.get())
                    room_charge = float(room_charge_entry.get())
                    
                    medicines_charges = float(medicines_charges_entry.get())
                    dressing_charges = float(dressing_charges_entry.get())
                    phoco_machine_charges = float(phoco_machine_charges_entry.get())
                    miscellaneous_charges = float(miscellaneous_charges_entry.get())
                    
                    
                    total_amount = (operative_fees + lens_implant + operation_theatre_charges + 
                                    anaesthetist_fees + viscoelastic + room_charge + medicines_charges + dressing_charges + 
                                    phoco_machine_charges + miscellaneous_charges)
                    
                    
                    total_amount_label.config(text=f"Total amount to be paid: {total_amount:.2f} Rs.")
                except ValueError:
                    
                    total_amount_label.config(text="Invalid input")

                
                

            def convert_to_words(num):
                if num == 0:
                    return "zero"
                
                ones = ["", "one", "two", "three", "four", "five", "six", "seven", "eight", "nine"]
                tens = ["", "", "twenty", "thirty", "forty", "fifty", "sixty", "seventy", "eighty", "ninety"]
                teens = ["ten", "eleven", "twelve", "thirteen", "fourteen", "fifteen", "sixteen", "seventeen", "eighteen", "nineteen"]
                
                words = ""
                
                # Handling lakhs
                if num >= 100000:
                    lakhs = num // 100000
                    words += ones[lakhs] + " lakh "
                    num %= 100000
                
                # Handling thousands
                if num >= 1000:
                    thousands = num // 1000
                    words += convert_to_words(thousands) + " thousand "
                    num %= 1000
                
                # Handling hundreds
                if num >= 100:
                    hundreds = num // 100
                    words += ones[hundreds] + " hundred "
                    num %= 100
                    
                # Handling tens and ones
                if num >= 10 and num <= 19:
                    words += teens[num - 10] + " "
                    num = 0
                elif num >= 20:
                    tens_digit = num // 10
                    words += tens[tens_digit] + " "
                    num %= 10
                    
                if num >= 1 and num <= 9:
                    words += ones[num] + " "
                    
                return words.strip()





            def generate_pdf_bill(patient_info, billing_details, doctor_info, total_amount, department, billing_items,adm_date,dis_date):
                c = canvas.Canvas("IPD_bill.pdf", pagesize=letter)
                
            
                bold_font = "Helvetica-Bold"
                normal_font = "Helvetica"
                title_font_size = 15
                section_font_size = 13
                content_font_size = 11
                
                
                c.setFont(bold_font, title_font_size)
                c.drawCentredString(300, 765, "EYE HOSPITAL, PUNE")
                
                # x1,y1,x2,y2
                c.line(50, 750, 550, 750)
                
                
                text_x = 50
                text_y = 720  
                
                line_height = 20
                
                
                c.setFont(bold_font, section_font_size)
                c.drawString(text_x, text_y, "Patient Information:")
                text_y -= line_height
                text_y = text_y-10

                for label, value in patient_info.items():
                    c.setFont(normal_font, content_font_size)
                    c.drawString(text_x, text_y, f"{label}:")
                    c.drawRightString(250, text_y, str(value))  # Aligning value to the right
                    text_y -= line_height
                
                text_y -= line_height  
                
            
                
                c.setFont(bold_font, section_font_size)
                c.drawString(350, 720, "Doctor's Information:")
                new_y=690
                for label, value in doctor_info.items():
                    c.setFont(normal_font, content_font_size)
                    c.drawString(350, new_y, f"{label}:")
                    c.drawRightString(560, new_y, str(value))  
                    new_y -= line_height
                
                
                
                new_y=new_y-10
                c.setFont(bold_font, section_font_size)
                c.drawString(350, new_y, "Department:")
                c.setFont(normal_font, content_font_size)
                c.drawString(540, new_y, department)  
                new_y=new_y-20
                c.drawString(350,new_y,f"Date Of Admission: {adm_date}")
                new_y=new_y-20
                c.drawString(350,new_y,f"Date Of Discharge: {dis_date}")    
                
                c.line(50, text_y+15, 550, text_y+15)
                text_y -= line_height
                
                c.setFont(bold_font, section_font_size)
                c.drawString(text_x, text_y, "Billing Items:")
                text_y -= line_height

                for key, value in billing_items.items():
                    c.setFont(normal_font, content_font_size)
                    c.drawString(text_x, text_y, f"{key}:")
                    c.drawRightString(250, text_y, str(value))  
                    text_y -= line_height
                        
                
                text_y -= line_height 
                
                
                c.setFont(bold_font, section_font_size)
                c.drawString(text_x, text_y, "Total Amount:")
                c.setFont(normal_font, content_font_size)
                c.drawRightString(250, text_y, f"Rs. {total_amount}")  

                c.setFont(bold_font, section_font_size)
                text_y -= line_height 
                c.line(50, text_y, 550, text_y)
                text_y -= line_height 
                text_y -= line_height
                c.drawString(text_x, text_y, "RECEIPT")
                text_y -= line_height
                c.setFont(normal_font, content_font_size)

                c.drawString(text_x, text_y,f"Received the sum of Rs {total_amount}")
                text_y=text_y-20

                amt=int(total_amount)
                result=convert_to_words(amt)

                c.drawString(text_x,text_y,f"In words : {result}")
                text_y=text_y-20
                c.drawString(text_x,text_y,"By CASH / CARD / CHEQUE / UPI")

                text_y=text_y-50
                c.setFont(bold_font,section_font_size)
                c.drawString(text_x+400,text_y,"Name of Doctor")
                c.setFont(normal_font,content_font_size)
                text_y=text_y-20
                c.drawString(text_x+430,text_y,"Degree")
            
                c.save()
                filename = "IPD_bill.pdf"
                
                if os.name == "posix": 
                    os.system("open " + filename)
                elif os.name == "nt": 
                    os.system("start " + filename)





            
            


            def print_receipt():
                
                patient_info = {
                    "Name": name,
                    "Age": str(doc[5]),
                    "Sex": str(doc[6]),
                    "Mobile": str(doc[8]),
                    "Date": today_string
                }
                adm_date=admission_date_entry.get()
                dis_date=discharge_date_entry.get()



                
                billing_items = {
                    "Diagnosis": diagnosis_entry.get(),
                    "Operative Procedure": operative_procedure_entry.get(),
                    "Operative Fees": operative_fees_entry.get(),
                    "Lens Implant": lens_implant_entry.get(),
                    "Operation Theatre Charges": operation_theatre_charges_entry.get(),
                    "Anaesthetist Fees": anaesthetist_fees_entry.get(),
                    "Viscoelastic": viscoelastic_entry.get(),
                    "Room charge": room_charge_entry.get(),
                    "Medicines charges": medicines_charges_entry.get(),
                    "Dressing charges": dressing_charges_entry.get(),
                    "Phoco Machine charges": phoco_machine_charges_entry.get(),
                    "Other Miscellaneous charges": miscellaneous_charges_entry.get()
                }

                
                doctor_info = {
                    "Doctor's Name": doctor_name_entry.get(),
                    "Doctor's Degree": doctor_degree_entry.get()
                }

                
                department = department_entry.get()

                
                # Extract only the numeric part of the string and remove any non-numeric characters
                total_amount_text = total_amount_label.cget("text").split(":")[-1].strip().replace('Rs.', '').strip()

                # Convert the extracted numeric part to a floating-point number
                total_amount_float = float(total_amount_text)

                # Convert the floating-point number to an integer (if needed) or use it directly
                total_amount = int(total_amount_float)  # Convert to integer if necessary




                
                generate_pdf_bill(patient_info, {}, doctor_info, total_amount, department,billing_items,adm_date,dis_date)



            
            
            
            
            
            
            patient_info_frame = ttk.LabelFrame(billing, text="Patient Information")
            patient_info_frame.grid(row=0, column=0, padx=50, pady=20, sticky="nsew")  

            ttk.Label(patient_info_frame, text="Bill No: ").grid(row=0, column=0, sticky="w")
            bill_no_entry = ttk.Entry(patient_info_frame)
            bill_no_entry.grid(row=0, column=0, sticky="e")


            
            
            ttk.Label(patient_info_frame, text="Name of Patient: " + name).grid(row=1, column=0, sticky="w")

            
            ttk.Label(patient_info_frame, text="Age: " + str(doc[5])).grid(row=2, column=0, sticky="w")

            
            ttk.Label(patient_info_frame, text="Sex: " + str(doc[6])).grid(row=3, column=0, sticky="w")

            
            ttk.Label(patient_info_frame, text="Mobile Number: " + str(doc[8])).grid(row=4, column=0, sticky="w")

            
            ttk.Label(patient_info_frame, text="Date: " + today_string).grid(row=5, column=0, sticky="w")

            department_label = ttk.Label(patient_info_frame, text="Department: ")
            department_label.grid(row=6, column=0, sticky="w")

            department_entry = ttk.Combobox(patient_info_frame, values=["OPD", "IPD"])
            department_entry.grid(row=6, column=1, sticky="w")

            
            admission_date_label = ttk.Label(patient_info_frame, text="Date of Admission: ")
            admission_date_label.grid(row=7, column=0, sticky="w")
            admission_date_entry = ttk.Entry(patient_info_frame)
            admission_date_entry.grid(row=7, column=1, sticky="w")

            admission_time_label = ttk.Label(patient_info_frame, text="Time of Admission: ")
            admission_time_label.grid(row=8, column=0, sticky="w")
            admission_time_entry = ttk.Entry(patient_info_frame)
            admission_time_entry.grid(row=8, column=1, sticky="w")

            discharge_date_label = ttk.Label(patient_info_frame, text="Date of Discharge: ")
            discharge_date_label.grid(row=9, column=0, sticky="w")
            discharge_date_entry = ttk.Entry(patient_info_frame)
            discharge_date_entry.grid(row=9, column=1, sticky="w")

            discharge_time_label = ttk.Label(patient_info_frame, text="Time of Discharge: ")
            discharge_time_label.grid(row=10, column=0, sticky="w")
            discharge_time_entry = ttk.Entry(patient_info_frame)
            discharge_time_entry.grid(row=10, column=1, sticky="w")







            
            billing_details_frame = ttk.LabelFrame(billing, text="Billing Details")
            billing_details_frame.grid(row=1, column=0, padx=50, pady=20, sticky="nsew") 

            

            diagnosis_label = ttk.Label(billing_details_frame, text="Diagnosis: ")
            diagnosis_label.grid(row=0, column=0, sticky="e")

            diagnosis_entry = ttk.Entry(billing_details_frame)
            diagnosis_entry.grid(row=0, column=1, sticky="w")

            operative_procedure_label = ttk.Label(billing_details_frame, text="Operative Procedure: ")
            operative_procedure_label.grid(row=1, column=0, sticky="e")
            operative_procedure_entry = ttk.Entry(billing_details_frame)
            operative_procedure_entry.grid(row=1, column=1, sticky="w")

            operative_fees_label = ttk.Label(billing_details_frame, text="Operative Fees: ")
            operative_fees_label.grid(row=2, column=0, sticky="e")
            operative_fees_entry = ttk.Entry(billing_details_frame)
            operative_fees_entry.grid(row=2, column=1, sticky="w")

            lens_implant_label = ttk.Label(billing_details_frame, text="Lens Implant: ")
            lens_implant_label.grid(row=3, column=0, sticky="e")
            lens_implant_entry = ttk.Entry(billing_details_frame)
            lens_implant_entry.grid(row=3, column=1, sticky="w")

            operation_theatre_charges_label = ttk.Label(billing_details_frame, text="Operation Theatre Charges: ")
            operation_theatre_charges_label.grid(row=4, column=0, sticky="e")
            operation_theatre_charges_entry = ttk.Entry(billing_details_frame)
            operation_theatre_charges_entry.grid(row=4, column=1, sticky="w")

            anaesthetist_fees_label = ttk.Label(billing_details_frame, text="Anaesthetist Fees: ")
            anaesthetist_fees_label.grid(row=5, column=0, sticky="e")
            anaesthetist_fees_entry = ttk.Entry(billing_details_frame)
            anaesthetist_fees_entry.grid(row=5, column=1, sticky="w")

            viscoelastic_label = ttk.Label(billing_details_frame, text="Viscoelastic: ")
            viscoelastic_label.grid(row=6, column=0, sticky="e")
            viscoelastic_entry = ttk.Entry(billing_details_frame)
            viscoelastic_entry.grid(row=6, column=1, sticky="w")

            room_charge_label = ttk.Label(billing_details_frame, text="Room Charge: ")
            room_charge_label.grid(row=7, column=0, sticky="e")
            room_charge_entry = ttk.Entry(billing_details_frame)
            room_charge_entry.grid(row=7, column=1, sticky="w")


            medicines_charges_label = ttk.Label(billing_details_frame, text="Injt + Medicines Charges: ")
            medicines_charges_label.grid(row=8, column=0, sticky="e")
            medicines_charges_entry = ttk.Entry(billing_details_frame)
            medicines_charges_entry.grid(row=8, column=1, sticky="w")

            dressing_charges_label = ttk.Label(billing_details_frame, text="Dressing Charges: ")
            dressing_charges_label.grid(row=9, column=0, sticky="e")
            dressing_charges_entry = ttk.Entry(billing_details_frame)
            dressing_charges_entry.grid(row=9, column=1, sticky="w")

            phoco_machine_charges_label = ttk.Label(billing_details_frame, text="Phoco Machine Charges: ")
            phoco_machine_charges_label.grid(row=10, column=0, sticky="e")
            phoco_machine_charges_entry = ttk.Entry(billing_details_frame)
            phoco_machine_charges_entry.grid(row=10, column=1, sticky="w")

            miscellaneous_charges_label = ttk.Label(billing_details_frame, text="Other Miscellaneous Charges: ")
            miscellaneous_charges_label.grid(row=11, column=0, sticky="e")
            miscellaneous_charges_entry = ttk.Entry(billing_details_frame)
            miscellaneous_charges_entry.grid(row=11, column=1, sticky="w")





            
            total_amount_label = ttk.Label(billing_details_frame, text="Total amount to be paid: ")
            total_amount_label.grid(row=12, column=0, sticky="e")

            calculate_amount_button = ttk.Button(billing_details_frame, text="Calculate Amount", command=calculate_amount)
            calculate_amount_button.grid(row=13, columnspan=2, pady=10)

            
            

            ttk.Button(billing, text="Print Receipt", command=print_receipt).grid(row=4, column=0, padx=50, pady=10, sticky="nsew")

            ttk.Button(billing, text="Exit Billing Section", command=billing.destroy).grid(row=5, column=0, padx=50, pady=10, sticky="nsew")

            



            doctor_details_frame = ttk.LabelFrame(billing, text="Doctor Details")
            doctor_details_frame.grid(row=0, column=1, padx=50, pady=20, sticky="nsew")



            
            doctor_name_label = ttk.Label(doctor_details_frame, text="Doctor's Name: ")
            doctor_name_label.grid(row=0, column=0, sticky="e")
            doctor_name_entry = ttk.Combobox(doctor_details_frame, values=["Dr. Pathan"])
            doctor_name_entry.grid(row=0, column=1, sticky="w")


            doctor_degree_label = ttk.Label(doctor_details_frame, text="Doctor's Degree: ")
            doctor_degree_label.grid(row=1, column=0, sticky="e")
            doctor_degree_entry = ttk.Combobox(doctor_details_frame, values=["MBBS", "MS", "MD"])
            doctor_degree_entry.grid(row=1, column=1, sticky="w")


       



            


        ipd_billing_button=ttk.Button(patient_detail_ipd,text="Billing",command=ipd_billing_section)
        # billing section func to be implemented
        ipd_billing_button.grid(row=13 , column= 3,sticky=tk.S)
        
        


        '''ipd billing ends'''
        
        
        
        
        def print_every():
            pdf = canvas.Canvas("hello.pdf")
            pdf.drawString(100, 800, name)
            pdf.drawString(100, 700, complaintxt.get("1.0",'end-1c'))
            pdf.drawString(100, 600, diagnosistxt.get("1.0",'end-1c'))
            pdf.drawString(100, 500, examtxt.get("1.0",'end-1c'))
            pdf.drawString(100, 400, medicinetxt.get("1.0",'end-1c'))
            pdf.drawString(100, 300, advisedtxt.get("1.0",'end-1c'))
            pdf.save()
            
            filename = "hello.pdf"
            
            if os.name == "posix":  # for macOS or Linux
                os.system("open " + filename)
            elif os.name == "nt":  # for Windows
                os.system("start " + filename)
        
        
        
        
        
root = Tk()
root.geometry("400x150")
root.title("Ophthalmic Software")


# Set the initial theme
root.tk.call("source", "azure.tcl")
root.tk.call("set_theme", "light")


usernameLabel = ttk.Label(root, text="User Name").grid(row=0, column=0)
username = StringVar()
usernameEntry = ttk.Entry(root, textvariable=username).grid(row=0, column=1)  

#password label and password entry box
passwordLabel = ttk.Label(root,text="Password").grid(row=1, column=0)  
password = StringVar()
passwordEntry = ttk.Entry(root, textvariable=password, show='*').grid(row=1, column=1)  

validateLogin = partial(validateLogin, username, password)
#login button
loginButton = ttk.Button(root, text="Login", command=validateLogin).grid(row=4, column=0)  


root.mainloop()





